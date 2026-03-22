import streamlit as st
from google import genai
import pandas as pd
import pyreadstat
import tempfile
import os
import re
import requests
from io import BytesIO
import plotly.express as px
import plotly.graph_objects as go
import scipy.stats as stats
from docx import Document
import random

# =================================================================
# 🔗 إعدادات قاعدة بيانات Firebase
# =================================================================
FIREBASE_DATABASE_URL = "https://bacflix-37fc3-default-rtdb.europe-west1.firebasedatabase.app"

# --- 1. إعدادات الواجهة الاحترافية (SPSS Style) ---
st.set_page_config(page_title="مختبر التحليل الديموغرافي والإحصائي", page_icon="📊", layout="wide")

st.markdown("""<style>
@import url('https://fonts.googleapis.com/css2?family=Cairo:wght@400;600;700&display=swap');
:root { --stat-blue: #0F62FE; --stat-dark: #121619; --stat-bg: #F4F7F6; --border-color: #E0E0E0; }
html, body, [class*="css"], .stMarkdown, p, h1, h2, h3, h4, h5, h6, label { font-family: 'Cairo', sans-serif; direction: rtl; text-align: right; }
.stApp { background-color: var(--stat-bg); }
.main .block-container { background: white; padding: 2rem; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.05); border: 1px solid var(--border-color); }
.stButton>button { background-color: var(--stat-blue); color: white; font-weight: 600; border-radius: 6px; border: none; padding: 0.6rem 1.2rem; width: 100%; transition: all 0.3s ease; }
.stButton>button:hover { background-color: #0353E9; box-shadow: 0 4px 10px rgba(15,98,254,0.3); }
.report-card { background: #ffffff; padding: 20px 25px; border-right: 5px solid var(--stat-blue); border-radius: 8px; margin: 15px 0; border: 1px solid var(--border-color); box-shadow: 0 2px 8px rgba(0,0,0,0.02); line-height: 1.9; color: #2c3e50; font-size: 1.05rem; }
[data-testid="stSidebar"] { background-color: var(--stat-dark) !important; color: white; }
[data-testid="stSidebar"] * { color: white !important; }
.creator-box { background: rgba(255,255,255,0.05); padding: 20px; border-radius: 10px; text-align: center; border: 1px solid rgba(255,255,255,0.1); margin-bottom: 20px; }
.creator-box h3 { color: #6EA6FF !important; margin-bottom: 5px; font-weight: 700;}
.creator-box p { margin: 2px 0; font-size: 0.9rem; opacity: 0.9;}
.info-box { background-color: #eef4ff; border-right: 4px solid #0F62FE; padding: 15px; border-radius: 5px; margin-bottom: 20px; color: #1a4e8a; font-weight: 600; }
</style>""", unsafe_allow_html=True)


# --- 2. إدارة المفاتيح السحابية (Firebase REST API) ---
def get_cloud_keys():
    try:
        response = requests.get(f"{FIREBASE_DATABASE_URL}/gemini_keys.json")
        if response.status_code == 200 and response.json():
            data = response.json()
            if isinstance(data, dict): return list(data.values())
            elif isinstance(data, list): return [k for k in data if k]
    except: pass
    return []

def save_key_to_cloud(new_key):
    try:
        existing_keys = get_cloud_keys()
        if new_key in existing_keys:
            return False, "⚠️ المفتاح موجود مسبقاً في قاعدة بيانات المنصة."
        requests.post(f"{FIREBASE_DATABASE_URL}/gemini_keys.json", json=new_key)
        return True, "✅ شكراً لك! تم دمج مفتاحك للعمل في المنصة."
    except Exception as e:
        return False, f"❌ خطأ في الاتصال بالسحابة: {e}"

def verify_gemini_key(test_key):
    try:
        client = genai.Client(api_key=test_key)
        client.models.generate_content(model="gemini-2.5-flash", contents="رد بكلمة 'ok' فقط.")
        return True, "✅ المفتاح صالح ومستعد للعمل."
    except Exception as e:
        err = str(e).lower()
        if "api_key_invalid" in err or "400" in err: return False, "❌ المفتاح وهمي أو غير صحيح."
        elif "quota" in err or "429" in err: return False, "⚠️ المفتاح صحيح لكنه استنفد الرصيد المجاني."
        else: return False, f"❌ خطأ غير متوقع: {e}"


# --- 3. محرك الاتصال الماسي ---
def call_gemini_sync(prompt):
    keys_to_try = []
    
    # 1. إعطاء الأولوية القصوى لمفتاح الباحث الحالي
    if st.session_state.get("user_api_key"):
        keys_to_try.append(st.session_state.user_api_key)
        
    # 2. جلب مفاتيح السيرفر (الخزنة + السحابة) كخطة بديلة لدعم الباحثين
    base_keys = list(st.secrets["API_KEYS"]) if "API_KEYS" in st.secrets else []
    cloud_keys = get_cloud_keys()
    server_keys = base_keys + cloud_keys
    random.shuffle(server_keys) 
    
    keys_to_try.extend(server_keys)
    
    if not keys_to_try:
        st.error("⚠️ لا توجد أي مفاتيح API متاحة للعمل.")
        return None
    
    for key in keys_to_try:
        try:
            client = genai.Client(api_key=key)
            response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
            return response.text
        except Exception:
            continue 
    return None


# --- 4. وظائف معالجة البيانات ---
class MockMeta:
    def __init__(self, columns):
        self.column_names = columns; self.column_labels = columns; self.variable_value_labels = {}

@st.cache_data
def load_data(file_bytes, file_name):
    ext = file_name.split('.')[-1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix="." + ext) as tmp:
        tmp.write(file_bytes); tmp_path = tmp.name
    try:
        if ext == 'sav': df, meta = pyreadstat.read_sav(tmp_path)
        elif ext == 'csv': df, meta = pd.read_csv(tmp_path), MockMeta(pd.read_csv(tmp_path).columns.tolist())
        else: df, meta = pd.read_excel(tmp_path), MockMeta(pd.read_excel(tmp_path).columns.tolist())
    finally: os.remove(tmp_path)
    return df, meta

def export_to_word(text):
    doc = Document()
    doc.add_heading('التقرير الديموغرافي والإحصائي', 0)
    doc.add_paragraph(re.sub(r'\*\*(.*?)\*\*', r'\1', text))
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def execute_safely(code, df, meta_dict):
    env = {'df': df.copy(), 'pd': pd, 'px': px, 'go': go, 'st': st, 'stats': stats, 'value_labels': meta_dict}
    try:
        exec(code, env)
        return True, None
    except Exception as e: return False, str(e)


# --- 5. القائمة الجانبية (Sidebar) ---
st.sidebar.markdown("""
<div class="creator-box">
    <h3>الدكتور قاسم سمير</h3>
    <p>تصميم وتطوير المنصة</p>
    <hr style="border-color: rgba(255,255,255,0.1); margin: 10px 0;">
    <p>📧 <a href="mailto:esa.gacem@univ-blida2.dz" style="color: #6EA6FF; text-decoration: none;">esa.gacem@univ-blida2.dz</a></p>
    <p>📞 0672595801</p>
</div>
""", unsafe_allow_html=True)

if st.sidebar.button("🧹 تصفير جلسة التحليل"):
    st.session_state.messages = []
    st.rerun()

# --- قسم مساهمة الباحثين والمفاتيح المشتركة ---
st.sidebar.markdown("---")
st.sidebar.markdown("### 🤝 مساهمة فريق البحث")
st.sidebar.info("لضمان استمرارية المنصة للفريق، نرجو إدراج مفتاحك الخاص. سيتم تفعيله لتجنب الضغط على سيرفيرات المنصة .")

st.sidebar.markdown("""
<a href="https://aistudio.google.com/app/apikey" target="_blank" style="text-decoration: none;">
    <div style="background-color: rgba(15, 98, 254, 0.1); border: 1px solid #0F62FE; padding: 10px; border-radius: 8px; text-align: center; margin-bottom: 15px; color: #6EA6FF; font-weight: bold; font-size: 0.9rem; transition: 0.3s;">
        🔗 احصل على مفتاحك مجاناً من هنا
    </div>
</a>
""", unsafe_allow_html=True)

if "user_api_key" not in st.session_state:
    st.session_state.user_api_key = None

user_key_input = st.sidebar.text_input("🔑 أدخل مفتاحك للمساهمة والتحليل:", type="password", placeholder="AIzaSy...")
if st.sidebar.button("➕ تفعيل ومساهمة"):
    if user_key_input:
        user_key = user_key_input.strip()
        with st.spinner("جاري فحص المفتاح..."):
            is_valid, msg = verify_gemini_key(user_key)
        
        if is_valid:
            # 1. تفعيل المفتاح للجلسة الحالية
            st.session_state.user_api_key = user_key
            st.sidebar.success("✅ تم تفعيل مفتاحك بنجاح لهذه الجلسة.")
            
            # 2. حفظ المفتاح في السحابة لدعم المجتمع
            saved, save_msg = save_key_to_cloud(user_key)
            if saved:
                st.sidebar.info(save_msg)
            else:
                st.sidebar.warning(save_msg) # إذا كان موجوداً مسبقاً
        else:
            st.sidebar.error(msg)
    else:
        st.sidebar.warning("الرجاء إدخال المفتاح أولاً.")


# --- 6. الواجهة الرئيسية (Main UI) ---
st.title("📊 مختبر التحليل الديموغرافي والإحصائي")
st.markdown("<p style='color: #666; font-size: 1.1rem; margin-bottom: 2rem;'>منصة استشارية متقدمة لتحليل البيانات الديموغرافية والاجتماعية للباحثين</p>", unsafe_allow_html=True)

if "df" not in st.session_state:
    st.session_state.df, st.session_state.meta, st.session_state.file_name = None, None, None
if "messages" not in st.session_state: st.session_state.messages = []

# رسالة ترحيبية وتوجيهية في الواجهة الرئيسية
if st.session_state.df is None:
    st.markdown("""
    <div class="info-box">
        💡 <b>ملاحظة للزملاء الباحثين:</b> لضمان أفضل تجربة وأسرع استجابة، نرجو إضافة مفتاح API الخاص بكم في القائمة الجانبية .
    </div>
    """, unsafe_allow_html=True)

uploaded_file = st.file_uploader("📂 قم برفع قاعدة البيانات للبدء (SPSS .sav, Excel .xlsx, CSV .csv)", type=['sav', 'csv', 'xlsx'])

if uploaded_file:
    if st.session_state.file_name != uploaded_file.name:
        with st.spinner("🔄 جاري قراءة وتشفير قاعدة البيانات..."):
            df, meta = load_data(uploaded_file.getvalue(), uploaded_file.name)
            st.session_state.df, st.session_state.meta, st.session_state.file_name = df, meta, uploaded_file.name
            st.session_state.messages = []
            st.rerun()

if st.session_state.df is not None:
    df, meta = st.session_state.df, st.session_state.meta
    meta_dict = dict(zip(meta.column_names, meta.column_labels)) if hasattr(meta, 'column_names') else {}
    
    c1, c2, c3 = st.columns(3)
    c1.metric("📌 إجمالي الحالات", f"{df.shape[0]:,}")
    c2.metric("📋 المتغيرات", df.shape[1])
    c3.metric("الملف النشط", st.session_state.file_name)

    with st.expander("🔍 استكشاف البيانات ودليل المتغيرات", expanded=False):
        t1, t2 = st.tabs(["البيانات الخام", "دليل الأعمدة (Labels)"])
        with t1: st.dataframe(df.head(10), use_container_width=True)
        with t2: 
            if meta_dict: st.dataframe(pd.DataFrame({"المتغير": list(meta_dict.keys()), "الوصف": list(meta_dict.values())}), use_container_width=True)
            else: st.info("لا توجد أوصاف برمجية.")

    st.markdown("---")

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]): st.markdown(msg["content"], unsafe_allow_html=True)

    if user_query := st.chat_input("✍️ اطلب تحليلاً (مثال: ادرس العلاقة بين المستوى التعليمي ومكان الإقامة)..."):
        st.session_state.messages.append({"role": "user", "content": user_query})
        with st.chat_message("user"): st.markdown(user_query)

        with st.chat_message("assistant"):
            with st.spinner("🔄 جاري التحليل وصياغة التقرير..."):
                prompt = f"""
                بناءً على الأعمدة والأوصاف المرفقة، قم بتحليل طبيعة البيانات أولاً:
                إذا كانت البيانات تحتوي على متغيرات ديموغرافية أو اجتماعية (مثل العمر، الجنس، الإقامة، التعليم، الدخل، الحالة العائلية...):
                - تصرف كأستاذ ديموغرافيا وخبير إحصائي رفيع المستوى.
                - استخرج الدلالات العميقة واستخدم مصطلحات ديموغرافية (مثل: التباين المجالي، التركيب النوعي، الفوارق السوسيو-اقتصادية، الدلالة الإحصائية).
                
                أما إذا كانت البيانات عامة ولا تتعلق بالديموغرافيا (مثل مبيعات، بيانات تقنية، مالية، أو غيرها):
                - تصرف كمحلل بيانات وخبير إحصائي محترف.
                - ركز على استخراج الاتجاهات، العلاقات، والإحصاءات الوصفية والاستدلالية بدقة دون استخدام مصطلحات ديموغرافية في غير محلها.

                البيانات متوفرة في المتغير `df`. 
                الأعمدة: {list(df.columns)}. 
                الأوصاف: {meta_dict}

                التعليمات الصارمة والمشتركة:
                1. **القراءة الأكاديمية/الاحترافية:** ابدأ دائماً بالشرح النظري وتحليل النتائج باللغة العربية بأسلوب علمي رصين قبل كتابة الكود.
                2. **توليد الكود:** ضع الكود البرمجي حصراً داخل علامات: ```python [الكود هنا] ```.
                3. **أدوات العرض:** لعرض الجداول استخدم `st.dataframe()`. لعرض الرسوم استخدم `st.plotly_chart(fig, use_container_width=True)` عبر مكتبة `px`.
                4. **المجاميع:** في الجداول المتقاطعة أو التكرارية، استخدم دائماً `margins=True, margins_name='المجموع'`.

                طلب المستخدم: {user_query}
                """
                
                response_text = call_gemini_sync(prompt)
                
                if response_text:
                    report_text = re.sub(r"```python\s*.*?```", "", response_text, flags=re.DOTALL|re.IGNORECASE).strip()
                    code_matches = re.findall(r"```python\s*(.*?)```", response_text, re.DOTALL|re.IGNORECASE)
                    
                    if report_text:
                        st.markdown(f'<div class="report-card">{report_text}</div>', unsafe_allow_html=True)
                        st.download_button("📄 تحميل التقرير (Word)", data=export_to_word(report_text), file_name='Analysis_Report.docx', mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document', key=f"dl_{len(st.session_state.messages)}")

                    if code_matches:
                        for code in code_matches:
                            success, error_msg = execute_safely(code.strip(), df, meta_dict)
                            if not success: st.error(f"⚠️ حدث خطأ في الرسم: {error_msg}")
                    
                    st.session_state.messages.append({"role": "assistant", "content": response_text})
                else:
                    st.error("❌ نعتذر، تعذر الاتصال بالمحركات. يرجى إدراج مفتاح API الخاص بك في القائمة الجانبية للمحاولة مجدداً.")
else:
    st.markdown("""
    <div style="text-align: center; padding: 3rem; background: white; border-radius: 10px; border: 1px dashed #ccc; margin-top: 2rem;">
        <h2 style="color: #0F62FE;">مرحباً بك في المختبر الديموغرافي</h2>
        <p style="color: #666; font-size: 1.1rem;">يرجى رفع قاعدة البيانات الخاصة بك في الأعلى للبدء.</p>
    </div>
    """, unsafe_allow_html=True)
